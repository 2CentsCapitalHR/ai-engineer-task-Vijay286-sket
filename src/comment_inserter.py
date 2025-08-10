import io
from typing import List
from docx import Document


def annotate_visible_notes(file_bytes: bytes, comments: List[str]) -> bytes:
    with io.BytesIO(file_bytes) as buffer:
        doc = Document(buffer)
    if comments:
        doc.add_page_break()
        doc.add_heading("Automated Review Notes", level=1)
        for c in comments:
            doc.add_paragraph(f"- {c}")
    out = io.BytesIO()
    doc.save(out)
    return out.getvalue()


